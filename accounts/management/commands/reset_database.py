from django.core.management.base import BaseCommand
from django.db import transaction
from django.contrib.auth import get_user_model
from accounts.models import Region, District, Sacco, ActivityLog
from members.models import Member, MemberGroup, MemberProfile
from loans.models import Loan, LoanProduct, LoanRepayment
from savings.models import SavingsAccount, SavingProduct, SavingsTransaction
from funding.models import Funding, FundingSource
from expenses.models import Expense, ExpenseCategory
from projects.models import Project
from notifications.models import Notification
import os
from pathlib import Path
import re

User = get_user_model()


class Command(BaseCommand):
    help = 'Reset database: Delete all data except system admin, then repopulate regions from Word documents'

    def add_arguments(self, parser):
        parser.add_argument(
            '--regions-dir',
            type=str,
            default='regions',
            help='Directory containing region DOCX files'
        )
        parser.add_argument(
            '--force',
            action='store_true',
            help='Force reset without confirmation'
        )

    def handle(self, *args, **options):
        regions_dir = options['regions_dir']
        force = options['force']
        
        if not force:
            confirm = input(
                self.style.WARNING(
                    '\n⚠️  WARNING: This will DELETE ALL DATA except system admin users!\n'
                    'This action cannot be undone. Type "YES" to continue: '
                )
            )
            if confirm != 'YES':
                self.stdout.write(self.style.ERROR('Operation cancelled.'))
                return
        
        # Get system admin users before deletion
        system_admins = list(User.objects.filter(is_system_admin=True).values(
            'id', 'username', 'email', 'first_name', 'last_name', 'is_system_admin',
            'is_staff', 'is_superuser', 'password'
        ))
        
        if not system_admins:
            self.stdout.write(self.style.ERROR('No system admin found. Aborting to prevent data loss.'))
            return
        
        self.stdout.write(self.style.WARNING(f'Found {len(system_admins)} system admin(s) to preserve'))
        
        with transaction.atomic():
            # Step 1: Delete all data (except system admins which are handled by CASCADE protection)
            self.stdout.write(self.style.WARNING('\n🗑️  Deleting all data...'))
            
            # Delete in reverse dependency order
            Notification.objects.all().delete()
            self.stdout.write('  ✓ Deleted notifications')
            
            LoanRepayment.objects.all().delete()
            self.stdout.write('  ✓ Deleted loan repayments')
            
            Loan.objects.all().delete()
            self.stdout.write('  ✓ Deleted loans')
            
            LoanProduct.objects.all().delete()
            self.stdout.write('  ✓ Deleted loan products')
            
            SavingsTransaction.objects.all().delete()
            self.stdout.write('  ✓ Deleted savings transactions')
            
            SavingsAccount.objects.all().delete()
            self.stdout.write('  ✓ Deleted savings accounts')
            
            SavingProduct.objects.all().delete()
            self.stdout.write('  ✓ Deleted saving products')
            
            Funding.objects.all().delete()
            self.stdout.write('  ✓ Deleted funding records')
            
            FundingSource.objects.all().delete()
            self.stdout.write('  ✓ Deleted funding sources')
            
            Expense.objects.all().delete()
            self.stdout.write('  ✓ Deleted expenses')
            
            ExpenseCategory.objects.all().delete()
            self.stdout.write('  ✓ Deleted expense categories')
            
            Project.objects.all().delete()
            self.stdout.write('  ✓ Deleted projects')
            
            MemberProfile.objects.all().delete()
            self.stdout.write('  ✓ Deleted member profiles')
            
            Member.objects.all().delete()
            self.stdout.write('  ✓ Deleted members')
            
            MemberGroup.objects.all().delete()
            self.stdout.write('  ✓ Deleted member groups')
            
            Sacco.objects.all().delete()
            self.stdout.write('  ✓ Deleted saccos')
            
            ActivityLog.objects.all().delete()
            self.stdout.write('  ✓ Deleted activity logs')
            
            District.objects.all().delete()
            self.stdout.write('  ✓ Deleted districts')
            
            Region.objects.all().delete()
            self.stdout.write('  ✓ Deleted regions')
            
            # Delete all users except system admins
            User.objects.exclude(is_system_admin=True).delete()
            self.stdout.write('  ✓ Deleted non-system-admin users')
            
            # Verify system admins still exist
            remaining_admins = User.objects.filter(is_system_admin=True).count()
            if remaining_admins != len(system_admins):
                self.stdout.write(self.style.ERROR(f'⚠️  Warning: Expected {len(system_admins)} system admins, found {remaining_admins}'))
            
            self.stdout.write(self.style.SUCCESS('\n✓ Database cleared successfully!'))
            
            # Step 2: Repopulate regions from Word documents
            self.stdout.write(self.style.SUCCESS('\n📂 Repopulating regions from Word documents...'))
            
            if not os.path.exists(regions_dir):
                self.stdout.write(self.style.ERROR(f'Regions directory not found: {regions_dir}'))
                return
            
            try:
                from docx import Document
            except ImportError:
                self.stdout.write(self.style.ERROR('python-docx is not installed. Please install it with: pip install python-docx'))
                return
            
            docx_files = list(Path(regions_dir).glob('*.docx'))
            
            if not docx_files:
                self.stdout.write(self.style.WARNING(f'No DOCX files found in {regions_dir}'))
                return
            
            self.stdout.write(self.style.SUCCESS(f'Found {len(docx_files)} DOCX files'))
            
            total_regions = 0
            total_districts = 0
            
            for docx_file in docx_files:
                try:
                    # Extract region name from filename (remove .docx extension)
                    region_name = docx_file.stem.strip()
                    # Clean up whitespace but preserve format
                    region_name = ' '.join(region_name.split())
                    
                    if not region_name:
                        self.stdout.write(self.style.WARNING(f'Skipping {docx_file.name}: Could not extract region name'))
                        continue
                    
                    # Create region
                    region = Region.objects.create(
                        name=region_name,
                        is_active=True
                    )
                    self.stdout.write(self.style.SUCCESS(f'✓ Created region: {region_name}'))
                    total_regions += 1
                    
                    # Parse districts from DOCX
                    districts = self.extract_districts_from_docx(docx_file)
                    
                    # Create districts
                    for district_name in districts:
                        district = District.objects.create(
                            name=district_name,
                            region=region,
                            is_active=True
                        )
                        self.stdout.write(f'  ✓ Created district: {district_name}')
                        total_districts += 1
                    
                    self.stdout.write(f'  Processed {len(districts)} districts for {region_name}\n')
                    
                except Exception as e:
                    self.stdout.write(self.style.ERROR(f'Error processing {docx_file.name}: {str(e)}'))
                    import traceback
                    traceback.print_exc()
                    continue
            
            self.stdout.write(self.style.SUCCESS(
                f'\n✅ Completed! Created {total_regions} regions and {total_districts} districts'
            ))
            self.stdout.write(self.style.SUCCESS(f'✓ Preserved {remaining_admins} system admin user(s)'))

    def extract_districts_from_docx(self, docx_file):
        """Extract district names from DOCX file"""
        try:
            from docx import Document
            doc = Document(str(docx_file))
        except Exception as e:
            self.stdout.write(self.style.WARNING(f'Could not read {docx_file.name}: {str(e)}'))
            return []
        
        districts = []
        district_keywords = ['district', 'districts', 'county', 'counties']
        
        # Try to find districts in paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if paragraph contains district keywords
            if any(keyword.lower() in text.lower() for keyword in district_keywords):
                # Try to extract district names from this paragraph
                extracted = self.extract_district_names_from_text(text)
                districts.extend(extracted)
            
            # Also check if the text looks like a district name
            if len(text) < 50 and text[0].isupper() and not any(char in text for char in [':', ';', '(', ')']):
                if text not in districts:
                    districts.append(text)
        
        # Try to extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) < 50:
                        if text[0].isupper() and text not in districts:
                            districts.append(text)
        
        # Clean and deduplicate districts
        cleaned_districts = []
        for district in districts:
            # Clean up the district name
            district = district.strip()
            district = re.sub(r'^\d+[\.\)]\s*', '', district)  # Remove numbering
            district = re.sub(r'^[-•]\s*', '', district)  # Remove bullets
            district = ' '.join(district.split())  # Normalize whitespace
            
            # Skip if too short or looks like a header
            if len(district) < 3:
                continue
            
            # Skip common non-district words
            skip_words = ['district', 'districts', 'region', 'regions', 'county', 'counties', 
                         'name', 'total', 'list', 'of', 'the', 'and', 'or', 'in', 'at', 'on',
                         'district name', 'district names', 'no.', 's/n', 'sn', 'number']
            if district.lower() in skip_words:
                continue
            
            # Skip if it contains only numbers or special chars
            if not re.search(r'[a-zA-Z]', district):
                continue
            
            # Skip if it's just a number
            if district.isdigit():
                continue
            
            if district and district not in cleaned_districts:
                cleaned_districts.append(district)
        
        return cleaned_districts

    def extract_district_names_from_text(self, text):
        """Extract district names from a text string"""
        districts = []
        
        # Pattern 1: "District: Name" or "Districts: Name1, Name2"
        match = re.search(r'districts?:\s*([^\n]+)', text, re.IGNORECASE)
        if match:
            district_list = match.group(1)
            # Split by comma or semicolon
            for item in re.split(r'[,;]', district_list):
                district = item.strip()
                if district and len(district) > 2:
                    districts.append(district)
        
        # Pattern 2: Look for capitalized words that might be district names
        words = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        for word in words:
            if len(word) > 3 and word.lower() not in ['the', 'and', 'for', 'are', 'with', 'from', 'this', 'that']:
                districts.append(word)
        
        return districts

