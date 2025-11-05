from django.core.management.base import BaseCommand
from django.contrib.auth import get_user_model
from django.db import transaction
from accounts.models import Region, District
from pathlib import Path
import os
import re

User = get_user_model()


class Command(BaseCommand):
    help = 'Setup system: create system admin and populate regions/districts from DOCX files'

    def add_arguments(self, parser):
        parser.add_argument('--username', type=str, default='system_admin', help='System admin username')
        parser.add_argument('--email', type=str, default='system@umsc.com', help='System admin email')
        parser.add_argument('--password', type=str, default='admin123', help='System admin password')
        parser.add_argument('--regions-dir', type=str, default='regions', help='Directory containing region DOCX files')
        parser.add_argument('--force', action='store_true', help='Proceed without interactive confirmation')

    def handle(self, *args, **options):
        username = options['username']
        email = options['email']
        password = options['password']
        regions_dir = options['regions_dir']
        force = options['force']

        if not force:
            confirm = input(self.style.WARNING(
                f"\nThis will ENSURE a system admin exists (username='{username}') and import regions/districts from '{regions_dir}'.\n"
                "Type 'YES' to continue: "
            ))
            if confirm != 'YES':
                self.stdout.write(self.style.ERROR('Operation cancelled.'))
                return

        with transaction.atomic():
            # 1) Ensure system admin exists
            admin, created = User.objects.get_or_create(
                username=username,
                defaults={
                    'email': email,
                    'is_system_admin': True,
                    'is_staff': True,
                    'is_superuser': True,
                    'first_name': 'System',
                    'last_name': 'Administrator',
                }
            )
            if created:
                admin.set_password(password)
                admin.save()
                self.stdout.write(self.style.SUCCESS(f"✓ Created system admin '{username}'"))
            else:
                # Update roles/flags to ensure elevated access
                updated = False
                if not admin.is_system_admin or not admin.is_staff or not admin.is_superuser:
                    admin.is_system_admin = True
                    admin.is_staff = True
                    admin.is_superuser = True
                    updated = True
                if email and admin.email != email:
                    admin.email = email
                    updated = True
                if updated:
                    admin.save()
                    self.stdout.write(self.style.SUCCESS(f"✓ Updated system admin '{username}' flags/email"))
                else:
                    self.stdout.write(f"System admin '{username}' already exists")

            # 2) Populate regions/districts from DOCX
            if not os.path.exists(regions_dir):
                self.stdout.write(self.style.ERROR(f"Regions directory not found: {regions_dir}"))
                return

            try:
                from docx import Document  # type: ignore
            except ImportError:
                self.stdout.write(self.style.ERROR("python-docx is not installed. Run: pip install python-docx"))
                return

            docx_files = list(Path(regions_dir).glob('*.docx'))
            if not docx_files:
                self.stdout.write(self.style.WARNING(f"No DOCX files found in {regions_dir}"))
                return

            self.stdout.write(self.style.SUCCESS(f"Found {len(docx_files)} DOCX files"))

            total_regions_created = 0
            total_districts_created = 0

            for docx_file in docx_files:
                try:
                    region_name = self._extract_region_name_from_filename(docx_file.stem)
                    if not region_name:
                        self.stdout.write(self.style.WARNING(f"Skipping {docx_file.name}: could not extract region name"))
                        continue

                    region, created_region = Region.objects.get_or_create(
                        name=region_name,
                        defaults={'is_active': True}
                    )
                    if created_region:
                        total_regions_created += 1
                        self.stdout.write(self.style.SUCCESS(f"✓ Created region: {region_name}"))
                    else:
                        self.stdout.write(f"Region already exists: {region_name}")

                    districts = self._extract_districts_from_docx(docx_file)

                    for district_name in districts:
                        _, created_district = District.objects.get_or_create(
                            name=district_name,
                            region=region,
                            defaults={'is_active': True}
                        )
                        if created_district:
                            total_districts_created += 1
                            self.stdout.write(f"  ✓ Created district: {district_name}")

                    self.stdout.write(f"  Processed {len(districts)} districts for {region_name}\n")

                except Exception as e:
                    self.stdout.write(self.style.ERROR(f"Error processing {docx_file.name}: {str(e)}"))
                    continue

            self.stdout.write(self.style.SUCCESS(
                f"\n✅ Done. Regions created: {total_regions_created}, Districts created: {total_districts_created}"
            ))

    def _extract_region_name_from_filename(self, filename: str) -> str:
        # Preserve filename formatting, normalize whitespace
        name = (filename or '').strip()
        return ' '.join(name.split()) if name else ''

    def _extract_districts_from_docx(self, docx_path: Path):
        try:
            from docx import Document  # type: ignore
            doc = Document(str(docx_path))
        except Exception:
            return []

        districts = []
        district_keywords = ['district', 'districts', 'county', 'counties']

        # Paragraphs
        for para in doc.paragraphs:
            text = (para.text or '').strip()
            if not text:
                continue
            if any(k in text.lower() for k in district_keywords):
                districts.extend(self._extract_district_names_from_text(text))
            if len(text) < 50 and text[:1].isupper() and not any(c in text for c in [':', ';', '(', ')']):
                if text not in districts:
                    districts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = (cell.text or '').strip()
                    if text and len(text) < 50 and text[:1].isupper() and text not in districts:
                        districts.append(text)

        # Clean
        cleaned = []
        for d in districts:
            d = d.strip()
            d = re.sub(r'^\d+[\.\)]\s*', '', d)
            d = re.sub(r'^[-•]\s*', '', d)
            d = ' '.join(d.split())
            if len(d) < 3:
                continue
            skip = {'district', 'districts', 'region', 'regions', 'county', 'counties',
                    'name', 'total', 'list', 'of', 'the', 'and', 'or', 'in', 'at', 'on',
                    'district name', 'district names', 'no.', 's/n', 'sn', 'number'}
            if d.lower() in skip:
                continue
            if not re.search(r'[a-zA-Z]', d):
                continue
            if d.isdigit():
                continue
            if d not in cleaned:
                cleaned.append(d)
        return cleaned

