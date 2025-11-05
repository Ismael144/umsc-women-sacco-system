from django.core.management.base import BaseCommand
from django.db import transaction
from accounts.models import Region, District
import os
from pathlib import Path
import re


class Command(BaseCommand):
    help = 'Populate regions and districts from DOCX files in the regions folder'

    def add_arguments(self, parser):
        parser.add_argument(
            '--regions-dir',
            type=str,
            default='regions',
            help='Directory containing region DOCX files'
        )

    def handle(self, *args, **options):
        regions_dir = options['regions_dir']
        
        if not os.path.exists(regions_dir):
            self.stdout.write(self.style.ERROR(f'Regions directory not found: {regions_dir}'))
            return

        try:
            from docx import Document
        except ImportError:
            self.stdout.write(self.style.ERROR('python-docx is not installed. Please install it with: pip install python-docx'))
            return

        docx_files = list(Path(regions_dir).glob('*.docx'))
        
        if not docx_files:
            self.stdout.write(self.style.WARNING(f'No DOCX files found in {regions_dir}'))
            return

        self.stdout.write(self.style.SUCCESS(f'Found {len(docx_files)} DOCX files'))
        
        total_regions = 0
        total_districts = 0
        
        with transaction.atomic():
            for docx_file in docx_files:
                try:
                    # Extract region name from filename
                    region_name = self.extract_region_name(docx_file.stem)
                    
                    if not region_name:
                        self.stdout.write(self.style.WARNING(f'Skipping {docx_file.name}: Could not extract region name'))
                        continue
                    
                    # Create or get region
                    region, created = Region.objects.get_or_create(
                        name=region_name,
                        defaults={'is_active': True}
                    )
                    
                    if created:
                        self.stdout.write(self.style.SUCCESS(f'Created region: {region_name}'))
                        total_regions += 1
                    else:
                        self.stdout.write(f'Region already exists: {region_name}')
                    
                    # Parse districts from DOCX
                    districts = self.extract_districts_from_docx(docx_file, region)
                    
                    # Create districts
                    for district_name in districts:
                        district, created = District.objects.get_or_create(
                            name=district_name,
                            region=region,
                            defaults={'is_active': True}
                        )
                        
                        if created:
                            self.stdout.write(f'  Created district: {district_name}')
                            total_districts += 1
                        else:
                            self.stdout.write(f'  District already exists: {district_name}')
                    
                    self.stdout.write(f'Processed {len(districts)} districts for {region_name}')
                    
                except Exception as e:
                    self.stdout.write(self.style.ERROR(f'Error processing {docx_file.name}: {str(e)}'))
                    continue
        
        self.stdout.write(self.style.SUCCESS(
            f'\nCompleted! Created {total_regions} new regions and {total_districts} new districts'
        ))

    def extract_region_name(self, filename):
        """Extract region name from filename, preserving the actual region name"""
        # Use the filename as-is, just clean up whitespace
        name = filename.strip()
        
        # Clean up whitespace but preserve the actual name format
        name = ' '.join(name.split())
        
        # Keep the original case/format from the filename
        # Don't change "ANKOLE KIGEZI REGION" to "Ankole Kigezi Region"
        # Just ensure proper spacing
        return name if name else None

    def extract_districts_from_docx(self, docx_file, region):
        """Extract district names from DOCX file"""
        try:
            from docx import Document
            doc = Document(str(docx_file))
        except Exception as e:
            self.stdout.write(self.style.WARNING(f'Could not read {docx_file.name}: {str(e)}'))
            return []
        
        districts = []
        district_keywords = ['district', 'districts', 'county', 'counties']
        
        # Try to find districts in paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Look for district names - could be in various formats
            # Common patterns:
            # - "District Name" or "District: Name"
            # - List items with district names
            # - Tables with district names
            
            # Check if paragraph contains district keywords
            if any(keyword.lower() in text.lower() for keyword in district_keywords):
                # Try to extract district names from this paragraph
                extracted = self.extract_district_names_from_text(text)
                districts.extend(extracted)
            
            # Also check if the text looks like a district name (simple heuristic)
            # District names are usually short, capitalized words
            if len(text) < 50 and text[0].isupper() and not any(char in text for char in [':', ';', '(', ')']):
                # Could be a district name
                if text not in districts:
                    districts.append(text)
        
        # Try to extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) < 50:
                        # Check if it looks like a district name
                        if text[0].isupper() and text not in districts:
                            districts.append(text)
        
        # Clean and deduplicate districts
        cleaned_districts = []
        for district in districts:
            # Clean up the district name
            district = district.strip()
            district = re.sub(r'^\d+[\.\)]\s*', '', district)  # Remove numbering
            district = re.sub(r'^[-•]\s*', '', district)  # Remove bullets
            district = ' '.join(district.split())  # Normalize whitespace
            
            # Skip if too short or looks like a header
            if len(district) < 3:
                continue
            
            # Skip common non-district words
            skip_words = ['district', 'districts', 'region', 'regions', 'county', 'counties', 
                         'name', 'total', 'list', 'of', 'the', 'and', 'or', 'in', 'at', 'on']
            if district.lower() in skip_words:
                continue
            
            # Skip if it contains only numbers or special chars
            if not re.search(r'[a-zA-Z]', district):
                continue
            
            if district and district not in cleaned_districts:
                cleaned_districts.append(district)
        
        return cleaned_districts

    def extract_district_names_from_text(self, text):
        """Extract district names from a text string"""
        districts = []
        
        # Common patterns:
        # - "District: Name" or "Districts: Name1, Name2"
        # - "Name District" format
        # - List items separated by commas, semicolons, or newlines
        
        # Pattern 1: "District: Name" or "Districts: Name1, Name2"
        match = re.search(r'districts?:\s*([^\n]+)', text, re.IGNORECASE)
        if match:
            district_list = match.group(1)
            # Split by comma or semicolon
            for item in re.split(r'[,;]', district_list):
                district = item.strip()
                if district and len(district) > 2:
                    districts.append(district)
        
        # Pattern 2: Look for capitalized words that might be district names
        # This is a simple heuristic - might need refinement
        words = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
        for word in words:
            if len(word) > 3 and word.lower() not in ['the', 'and', 'for', 'are', 'with']:
                districts.append(word)
        
        return districts
